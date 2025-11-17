"""
Create Word document from the user manual markdown.
Converts USER_MANUAL.md to a Word document.
"""
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re
    
    def create_word_manual():
        """Create Word document from markdown."""
        print("Creating Word document from USER_MANUAL.md...")
        
        # Read markdown file
        with open("USER_MANUAL.md", "r", encoding="utf-8") as f:
            content = f.read()
        
        # Create Word document
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Split into lines
        lines = content.split('\n')
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # Skip empty lines
            if not line:
                i += 1
                continue
            
            # Title (first line)
            if i == 0:
                title = doc.add_heading(line, level=0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                i += 1
                continue
            
            # Heading level 1
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
                i += 1
                continue
            
            # Heading level 2
            if line.startswith('## '):
                doc.add_heading(line[3:], level=2)
                i += 1
                continue
            
            # Heading level 3
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
                i += 1
                continue
            
            # Code blocks
            if line.startswith('```'):
                i += 1
                code_lines = []
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                if code_lines:
                    para = doc.add_paragraph('\n'.join(code_lines))
                    para.style = 'Intense Quote'
                i += 1
                continue
            
            # Bullet points
            if line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
                i += 1
                continue
            
            # Numbered list
            if re.match(r'^\d+\.\s', line):
                doc.add_paragraph(re.sub(r'^\d+\.\s', '', line), style='List Number')
                i += 1
                continue
            
            # Tables
            if '|' in line and line.count('|') >= 2:
                # Skip separator lines
                if '---' in line:
                    i += 1
                    continue
                
                # Parse table row
                cells = [c.strip() for c in line.split('|') if c.strip()]
                if cells:
                    if not hasattr(doc, '_current_table') or doc._current_table is None:
                        doc._current_table = doc.add_table(rows=1, cols=len(cells))
                        doc._current_table.style = 'Light Grid Accent 1'
                        row = doc._current_table.rows[0]
                        for j, cell_text in enumerate(cells):
                            row.cells[j].text = cell_text
                            # Make header bold
                            for paragraph in row.cells[j].paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                    else:
                        row = doc._current_table.add_row()
                        for j, cell_text in enumerate(cells):
                            row.cells[j].text = cell_text
                i += 1
                continue
            
            # Regular paragraph
            # Handle inline code
            para = doc.add_paragraph()
            parts = re.split(r'(`[^`]+`)', line)
            for part in parts:
                if part.startswith('`') and part.endswith('`'):
                    run = para.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                else:
                    para.add_run(part)
            
            i += 1
        
        # Save document
        output_file = "Dia_Model_User_Manual.docx"
        doc.save(output_file)
        print(f"\n[SUCCESS] Word document created: {output_file}")
        print(f"File saved in current directory.")
        print(f"\nYou can now:")
        print(f"  1. Open {output_file} in Microsoft Word")
        print(f"  2. Edit and format as needed")
        print(f"  3. Save and share with users")
        
except ImportError:
    print("python-docx not installed.")
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx"])
    
    print("\nPlease run this script again:")
    print("python create_word_manual.py")

if __name__ == "__main__":
    create_word_manual()

